# -*-coding: utf-8 -*-

import pandas as pd
from docx import Document
import platform
import datetime
import subprocess as sp
from datetime import timedelta
from tabulate import tabulate
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText


def sys_check():
    if "Darwin" in platform.system():
        file_edit_mac()
    elif "Windows" in platform.system():
        file_edit_win()
    elif "linux" in platform.system():
        file_edit_linux()
    else:
        print(platform.system() + "nie jest wspieraną platformą.")


def raport_doc():
    document = Document()
    document.add_heading("Raport z dnia "+ date, 0)
    document.add_paragraph("Cześć!")
    d = document.add_paragraph("Oto stan magazynowy z dnia " + date + ". ")
    d.add_run("Na ten moment nie ma koniecznośći zamawiania dostawy, ")
    d.add_run("ale musimy zamówić kuriera w celu odwesłania zwrotów.")
    d.add_run(" ")
    d.add_run("Pozdrawiam,")
    table = document.add_table(rows=(df1.shape[0]), cols=df1.shape[1])  # First row are table headers!
    for i, column in enumerate(df1) :
        for row in range(df1.shape[0]) :
            table.cell(row, i).text = str(df1[column][row])


    document.add_page_break()

    document.save('demo.docx')



def clear_mail():
    with open('stan.txt', 'r+') as f:
        t = f.read()
        to_delete = yesterday.strip()
        f.seek(0)
        for line in t.split('\n'):
            if line != to_delete:
                f.write(line + '\n')
        f.truncate()

def file_edit_win():
    programName = "notepad.exe"
    textfile = "stan.txt"
    p = sp.Popen([programName, textfile])
    p.wait()


def file_edit_mac():
    textfile = "demo.docx"
    p = sp.Popen(["open", "-W", textfile])
    p.wait()



def file_edit_linux():
    textfile = "stan.txt"
    p = sp.Popen(["gedit", textfile])
    p.wait()


#sys_check()


now = datetime.datetime.now() - timedelta(1)
yesterday = "9.11.2018"
#yesterday = now.strftime("%d.%m.%Y")
date = "10.11.2018"
df = pd.read_excel('stan.xlsx')
df = df.fillna(0)
df1 = df.loc[df["DATA"] == date]
output = df1.to_dict(orient='list')
table = tabulate(output, tablefmt = "grid", headers = "keys")

#raport_doc()

#clear_mail()

stan = open("stan.txt", "a")
stan.write("\n")
stan.write("\n")
stan.write(table)
stan.close()
print(df1)
